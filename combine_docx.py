from docx import Document
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt
from win32com import client as wc
from docx.oxml.ns import qn
import re

path = os.getcwd()
for i in os.listdir("files"):
    doc_all = Document()
    doc_all.styles['Normal'].font.size = Pt(10.5)
    doc_all.styles['Normal'].font.name = u'仿宋_GB2312'
    doc_all.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    for j in os.listdir("files/" + i):
        print("files/" + i + '/' + j)
        filename, aff = os.path.splitext(j)
        if aff == ".doc":
            w = wc.DispatchEx('Word.Application')
            doc = w.Documents.Open(path + "/files/" + i + '/' + j)
            doc.SaveAs(path + "/files/" + i + '/' + filename + ".docx", 12)
            doc.Close()
    for j in os.listdir("files/" + i):
        filename, aff = os.path.splitext(j)
        if aff == ".docx":
            doc = Document("files/" + i + '/' + j)
            print("files/" + i + '/' + j)
            count = 0
            for para in doc.paragraphs:
                temp = para.text
                if len(temp) == 0 or temp.isspace():
                    continue
                count += 1
                if count == 1 and len(temp) < 20:
                    paragraph = doc_all.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run(temp.replace(" ", ""))
                    run.font.name = u'方正小标宋简体'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
                elif (count == 2 or count == 3) and len(temp) < 40:
                    if count == 2:
                        if "■" not in temp:
                            x = re.split(r"\s+", temp)
                            if len(x) < 2:
                                continue
                            temp = " ".join(x[0:-1]) + " ■" + x[-1]
                        else:
                            temp = re.sub(r"\s*■", " ■", temp)
                    if count == 3 and "□" not in temp:
                        temp = "□" + temp
                    paragraph = doc_all.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run(temp)
                    run.font.name = u'楷体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
                else:
                    paragraph = doc_all.add_paragraph(temp.strip())
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    paragraph_format.first_line_indent = Pt(21)
            doc_all.add_page_break()
    doc_all.save(i + '.docx')
